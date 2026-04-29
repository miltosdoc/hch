from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('Integration Description Questionnaire', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Webdoc API – Hjärtcentrum Halland')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

def heading(n, t):
    doc.add_heading(f'{n}. {t}', level=1)

def body(t):
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(6)
    return p

def table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
        for p in tbl.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for rd in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(rd):
            cells[i].text = v
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# =================================================================
# 1
# =================================================================
heading('1', 'Background')

body(
    'Vi behöver automatisera hanteringen av inkommande remisser samt '
    'ta fram väntetidsstatistik inför en extern granskning. '
    'Remisser skannas och laddas upp i patientakter, varpå en osignerad '
    'anteckning skapas för personalens verifiering i signeringskorgen. '
    'Parallellt behöver vi extrahera registreringsdatum och datum för '
    'betalningsförbindelser ur bokningsdata, då dessa fält saknas i '
    'Webdocs CSV-exportmallar.'
)

# =================================================================
# 2
# =================================================================
heading('2', 'Description of the Solution')

body('Integrationen har två arbetsflöden:')

doc.add_heading('Arbetsflöde 1 – Dokumentuppladdning med verifiering', level=2)

steps1 = [
    'Autentisering (POST /oauth/token).',
    'Hämta kliniker, dokumenttyper, bokningstyper och användare.',
    'Extrahera personnummer från filnamn, slå upp patient (GET /v2/patients).',
    'Ladda upp dokument (POST Documents) med documentTypeId "Inreferral".',
    'Skapa osignerad anteckning (POST /v1/notes) kopplad till patienten.',
    'Lägg till journaltext i anteckningen (PATCH Records) med information om uppladdad remiss.',
    'Personalen verifierar och signerar anteckningen via signeringskorgen i Webdoc.',
]
for s in steps1:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()
doc.add_heading('Arbetsflöde 2 – Statistikuttag', level=2)

steps2 = [
    'Autentisering (POST /oauth/token).',
    'Hämta bokningar (GET Bookings) filtrerat på datumintervall.',
    'Extrahera registreringsdatum (bokningsdatum), patienttyp och betalningsdata '
    '(bookedPatientType, payments, freeCard med validFrom/validUntil).',
    'Komplettera med patienttyper (GET /v1/patientTypes), åtgärdskoder (GET /v1/actionCodes) '
    'och besöksdata (GET Visits) vid behov.',
    'Exportera till CSV/Excel för rapportering.',
]
for s in steps2:
    doc.add_paragraph(s, style='List Number')

# =================================================================
# 3
# =================================================================
heading('3', 'Webdoc API Endpoints')

table(
    ['Endpoint', 'Method', 'Beskrivning'],
    [
        ['/oauth/token', 'POST', 'Autentisering'],
        ['/v1/clinics', 'GET', 'Kliniker'],
        ['/v1/documentTypes', 'GET', 'Dokumenttyper'],
        ['/v1/clinics/{clinicId}/documents', 'POST', 'Ladda upp dokument'],
        ['/v2/patients', 'GET', 'Patientuppslag'],
        ['/v1/patientTypes', 'GET', 'Patienttyper'],
        ['/v1/patientTypes/{id}', 'GET', 'Patienttyp per ID'],
        ['/v1/bookings', 'GET', 'Bokningsdata'],
        ['/v1/bookings', 'POST', 'Skapa bokning'],
        ['/v1/bookings', 'PATCH', 'Uppdatera bokning'],
        ['/v1/bookings', 'DEL', 'Ta bort bokning'],
        ['/v1/bookingTypes', 'GET', 'Bokningstyper'],
        ['/v1/notes', 'POST', 'Skapa anteckning'],
        ['/v1/visits/{visitId}/records', 'PATCH', 'Journaldata till anteckning/besök'],
        ['/v1/actionCodes', 'GET', 'Åtgärdskoder (KVÅ)'],
        ['/v1/users', 'GET', 'Användare/vårdgivare'],
        ['/v1/clinics/{clinicId}/visits', 'GET', 'Besöksdata'],
        ['/v1/visits/{visitId}/recordSignatures', 'POST', 'Signera journalanteckning'],
    ]
)

# =================================================================
# 4
# =================================================================
heading('4', 'Webdoc API Scopes')

table(
    ['Scope', 'Beskrivning'],
    [
        ['documents:write', 'Uppladdning av dokument'],
        ['document-types:read', 'Läsa dokumenttyper'],
        ['clinics:read', 'Läsa klinikinformation'],
        ['patient:read', 'Läsa patientdata'],
        ['patient:write', 'Skapa anteckningar och journaldata'],
        ['patient-types:read', 'Läsa patienttyper'],
        ['bookings:read', 'Läsa bokningsdata'],
        ['bookings:write', 'Hantera bokningar'],
        ['actioncodes:read', 'Läsa åtgärdskoder'],
        ['users:read', 'Läsa användardata'],
        ['record-signatures:write', 'Hantera signaturer på journalanteckningar'],
        ['self-service', 'Läsa besöksdata (GET /v1/clinics/{clinicId}/visits)'],
    ]
)

# =================================================================
# 5
# =================================================================
heading('5', 'Webdoc API Grant Types')

p = doc.add_paragraph()
r = p.add_run('Grant type: ')
r.bold = True
p.add_run('client_credentials')
body('Server-till-server-integration. Ingen redirect URI.')

# =================================================================
# 6
# =================================================================
heading('6', 'External Links ("uthopp")')
body('Inga. Backend-integration utan gränssnitt i Webdoc.')

# =================================================================
# 7
# =================================================================
heading('7', 'Risks')

table(
    ['Risk', 'Åtgärd'],
    [
        ['Max 100 bokningar per anrop', 'Paginering via offset-parameter'],
        ['Felaktigt personnummer i filnamn', 'Formatvalidering före uppladdning'],
        ['API-driftstörning', 'Retry-logik och loggning'],
        ['Känsliga personuppgifter', 'Ingen permanent lagring; GDPR-rutiner följs'],
        ['Token-utgång under batch', 'Automatisk token-förnyelse'],
    ]
)

# =================================================================
# 8
# =================================================================
heading('8', 'Testing Process')

body(
    'Fullständig testning i integrationsmiljön '
    '(api-integration.carasent.net / auth-integration.carasent.net). '
    'Alla nya endpoints testas i integrationsmiljön innan produktionsanvändning.'
)

# =================================================================
# 9
# =================================================================
heading('9', 'Release Process')

body(
    'Lokala Python-skript på säkra arbetsstationer. Uppdateringar testas i '
    'integrationsmiljön, driftsätts manuellt efter godkänd verifiering.'
)

# =================================================================
# 10
# =================================================================
heading('10', 'Technical Contact')

t = doc.add_table(rows=4, cols=2)
t.style = 'Light Grid Accent 1'
for i, (l, v) in enumerate([
    ('Name:', 'Miltiadis Triantafyllou'),
    ('Position/title:', 'CEO'),
    ('Email:', 'miltiadis@hjartcentrumhalland.se'),
    ('Phone:', '0720233343'),
]):
    t.rows[i].cells[0].text = l
    t.rows[i].cells[1].text = v
    for p in t.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

doc.add_paragraph()

# =================================================================
# 11
# =================================================================
heading('11', 'Support Contact')

t2 = doc.add_table(rows=3, cols=2)
t2.style = 'Light Grid Accent 1'
for i, (l, v) in enumerate([
    ('Email:', 'info@hjartcentrumhalland.se'),
    ('Name:', 'Miltiadis Triantafyllou'),
    ('Phone:', '010-300 19 20'),
]):
    t2.rows[i].cells[0].text = l
    t2.rows[i].cells[1].text = v
    for p in t2.rows[i].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

# ---- Save ----
output_path = r'C:\Users\Miltiadis.t\Desktop\WebdocAPI\Integration_Questionnaire_Webdoc_Hjartcentrum_Halland.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
